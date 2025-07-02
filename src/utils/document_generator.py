import os
from docx import Document

class DocumentGenerator:
    def __init__(self, modelo_path: str):
        if not os.path.exists(modelo_path):
            raise FileNotFoundError(f"Modelo não encontrado em: {modelo_path}")
        self.modelo_path = modelo_path

    def gerar_contestacao_word(self, dados_extraidos: dict, corpo_ia: str, dados_advogado: dict, output_path: str) -> str:
        """
        Gera uma contestação preenchendo os placeholders em um modelo .docx.

        Substitui campos como:
        - [NOME_AUTOR], [NOME_REU], [TIPO_ACAO], [VALOR_CAUSA]
        - [NOME_ADVOGADO], [OAB_ADVOGADO], [UF_ADVOGADO]
        - {{corpo_contestacao}} (texto principal gerado pela IA)
        """
        doc = Document(self.modelo_path)

        substituicoes = {
            "[NOME_AUTOR]": dados_extraidos.get("autor", ""),
            "[NOME_REU]": dados_extraidos.get("reu", ""),
            "[TIPO_ACAO]": dados_extraidos.get("tipo_acao", ""),
            "[VALOR_CAUSA]": dados_extraidos.get("valor_causa", ""),
            "[NOME_ADVOGADO]": dados_advogado.get("nome_advogado", ""),
            "[OAB_ADVOGADO]": dados_advogado.get("oab_advogado", ""),
            "[UF_ADVOGADO]": dados_advogado.get("uf_advogado", ""),
            "{{corpo_contestacao}}": corpo_ia.strip()
        }

        def substituir_em_paragrafos(paragrafos):
            for p in paragrafos:
                for key, val in substituicoes.items():
                    if key in p.text:
                        for run in p.runs:
                            run.text = run.text.replace(key, val)

        # Substituição no corpo
        substituir_em_paragrafos(doc.paragraphs)

        # Substituição em tabelas (caso o modelo contenha)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    substituir_em_paragrafos(cell.paragraphs)

        doc.save(output_path)
        return output_path


